#!/usr/bin/env python3
"""Verify that every tokenized template contains real {{TOKEN}} placeholders.
Also enforce the OPSEC rule for sub-facing docs."""

import os, re, sys
from docx import Document
import openpyxl

ROOT = os.path.join(os.path.dirname(__file__), '..', '..', 'src', 'lib', 'bid-templates')
TOKENIZED = os.path.join(ROOT, 'tokenized')
TOKEN_RE = re.compile(r'\{\{[A-Z_0-9]+\}\}')
NAME_RE  = re.compile(r'\{\{([A-Z_0-9]+)\}\}')

OPSEC_FORBIDDEN = {
    'SOL_NUMBER', 'SOL_AGENCY', 'SOL_SUB_AGENCY', 'SOL_NOTICE_ID',
    'SOL_SAM_URL', 'SOL_ESTIMATED_VALUE', 'SOL_INCUMBENT_CONTRACTOR',
    'SOL_AGENCY_ABBREV', 'SOL_FACILITY_NAME',
    'SOL_SITE_1', 'SOL_SITE_2', 'SOL_SITE_3',
}

SUB_FACING = {
    '05_Subcontractor_Scope_of_Work.docx',
    '06_Sub_Outreach_Email.docx',
    '14_SIF_Subcontractor_Information_Form.docx',
}

def extract_text(path):
    if path.endswith('.docx'):
        d = Document(path)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for cp in cell.paragraphs:
                        parts.append(cp.text)
        for sec in d.sections:
            for hf in [sec.header, sec.footer]:
                for p in hf.paragraphs:
                    parts.append(p.text)
        return ' '.join(parts)
    if path.endswith('.xlsx'):
        wb = openpyxl.load_workbook(path)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for v in row:
                    if v is not None:
                        parts.append(str(v))
        return ' '.join(parts)
    return ''

def main():
    if not os.path.isdir(TOKENIZED):
        print(f'ERROR: {TOKENIZED} does not exist.')
        sys.exit(2)
    failed = []
    for f in sorted(os.listdir(TOKENIZED)):
        if not f.endswith(('.docx', '.xlsx')):
            continue
        text = extract_text(os.path.join(TOKENIZED, f))
        tokens = TOKEN_RE.findall(text)
        names = {m.group(1) for m in NAME_RE.finditer(text)}
        if len(tokens) == 0:
            failed.append(f'{f}: NO TOKENS')
            print(f'FAIL  {f}: 0 tokens')
            continue
        if f in SUB_FACING:
            forbid = (names & OPSEC_FORBIDDEN) | {n for n in names if n.startswith('CO_')}
            if forbid:
                failed.append(f'{f}: OPSEC leak — {sorted(forbid)}')
                print(f'FAIL  {f}: OPSEC leak {sorted(forbid)}')
                continue
        print(f'OK    {f}: {len(tokens)} tokens, {len(names)} unique')
    if failed:
        print(f'\n{len(failed)} failures.')
        sys.exit(1)
    print('\nAll templates passed.')

if __name__ == '__main__':
    main()
