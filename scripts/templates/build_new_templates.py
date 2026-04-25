#!/usr/bin/env python3
"""
Build the 7 templates Ella does NOT have samples for:
  03_Wage_Determination.docx
  04b_Proposal_Technical.docx
  06_Sub_Outreach_Email.docx
  08_Tracker_Workbook.xlsx (Excel, multi-sheet)
  09_To_Do_List.docx
  10_Risk_Log.docx
  12_Game_Plan.docx

Also rename 04_Proposal_Pricing.docx -> 04a_Proposal_Pricing.docx in tokenized/.
And clean up Phase-3 files (15-19) from tokenized/ since they belong to Phase 3.
"""

import os, shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

ROOT = '/sessions/beautiful-lucid-ritchie/mnt/mog-tracker-app/src/lib/bid-templates'
OUT = os.path.join(ROOT, 'tokenized')

# 1. Rename 04 -> 04a (tokenized only; source stays as 04 since Ella's filename)
src_04 = os.path.join(OUT, '04_Proposal_Pricing.docx')
dst_04a = os.path.join(OUT, '04a_Proposal_Pricing.docx')
if os.path.exists(src_04) and not os.path.exists(dst_04a):
    shutil.move(src_04, dst_04a)
    print('Renamed 04 -> 04a')

# 2. Clean up Phase-3 files from tokenized/ (they were left as zero-token copies)
for f in ['15_Cover_Letter_Submission.docx', '16_Contractor_Information_Sheet.docx',
          '17_Pricing_Schedule.docx', '18_Subcontractor_Data_Worksheet.docx',
          '19_Representations_and_Certifications.docx']:
    p = os.path.join(OUT, f)
    if os.path.exists(p):
        try:
            os.remove(p)
        except PermissionError:
            print(f"  (skipped, permission: {f})")
        print(f'Removed Phase-3 file from tokenized/: {f}')

# Also remove obsolete _processing_log.json if present
for f in ['_processing_log.json']:
    p = os.path.join(OUT, f)
    if os.path.exists(p):
        try:
            os.remove(p)
        except PermissionError:
            print(f"  (skipped, permission: {f})")


# ─── Helpers ─────────────────────────────────────────────────────────────────

def add_title(doc, text, size=18, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    p.style = doc.styles['Heading 2']


def add_h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    p.style = doc.styles['Heading 3']


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    r.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_kv_table(doc, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    if widths is None:
        widths = (Inches(2.0), Inches(4.5))
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        # Bold the label
        for p in table.cell(i, 0).paragraphs:
            for r in p.runs:
                r.bold = True
        try:
            table.columns[0].width = widths[0]
            table.columns[1].width = widths[1]
        except Exception:
            pass
    return table


# ─── 03_Wage_Determination.docx ──────────────────────────────────────────────

def build_03():
    doc = Document()
    add_title(doc, 'WAGE DETERMINATION REFERENCE')
    add_p(doc, 'Solicitation: {{SOL_NUMBER}} — {{SOL_TITLE}}', bold=True)
    add_p(doc, 'Agency: {{SOL_AGENCY}} ({{SOL_SUB_AGENCY}})')
    add_p(doc, 'Place of Performance: {{SOL_PLACE_OF_PERFORMANCE}}')
    add_p(doc, 'Compiled: {{TODAY_DATE}}')
    doc.add_paragraph()

    add_h2(doc, '1. Applicable Wage Determination')
    add_kv_table(doc, [
        ('Wage Determination Number:', '{{WAGE_DETERMINATION_NUMBER}}'),
        ('Revision Date:', '{{WAGE_DETERMINATION_REVISION_DATE}}'),
        ('Locality:', '{{WAGE_DETERMINATION_LOCATION}}'),
        ('Statute:', 'Service Contract Act of 1965 (41 U.S.C. § 6701 et seq.)'),
        ('Source:', 'sam.gov / wdol.gov as published by the U.S. Department of Labor'),
    ])

    doc.add_paragraph()
    add_h2(doc, '2. Mandatory Hourly Rates and Fringe Benefits')
    add_p(doc, 'The following rates are the legal floor for this contract. No bid may pay less than these amounts. Pay above the floor is permitted; pay below is a violation of the SCA and grounds for contract termination, debarment, and back-wage liability.')
    doc.add_paragraph()
    add_p(doc, '{{WAGE_RATES_TABLE}}')
    doc.add_paragraph()
    add_p(doc, 'Health and welfare fringe rate: {{WAGE_FRINGE}} per hour (or {{WAGE_FRINGE_EO13706}} for contracts subject to EO 13706).')
    add_p(doc, 'Vacation, holiday, and other fringe components are calculated per the wage determination cover sheet.')

    add_h2(doc, '3. Compliance Statement')
    add_p(doc, '{{ENTITY_NAME}} certifies that all labor performed under any resulting contract will be paid at or above the rates set forth in {{WAGE_DETERMINATION_NUMBER}}, including the required health and welfare fringe benefit rate. Certified payroll records will be maintained and produced upon request by the Contracting Officer or Department of Labor.')

    add_h2(doc, '4. Annual Adjustments')
    add_p(doc, 'SCA wage determinations are revised annually. {{ENTITY_NAME}} reserves the right to request an Equitable Adjustment under FAR 52.222-43 (Fair Labor Standards Act and Service Contract Labor Standards — Price Adjustment) for any option year in which the applicable wage determination is revised upward by the U.S. Department of Labor.')

    add_h2(doc, '5. References')
    add_bullet(doc, 'FAR 52.222-41 — Service Contract Labor Standards')
    add_bullet(doc, 'FAR 52.222-43 — FLSA and SCA Price Adjustment')
    add_bullet(doc, '29 CFR Part 4 — Labor Standards for Federal Service Contracts')
    add_bullet(doc, 'Executive Order 13706 — Establishing Paid Sick Leave for Federal Contractors')
    add_bullet(doc, 'sam.gov Wage Determinations: https://sam.gov/wage-determinations')

    add_h2(doc, '6. Document Verification')
    add_p(doc, 'Wage determination data above was extracted from the solicitation package and verified against the published U.S. Department of Labor wage determination on {{TODAY_DATE}}. The PDF copy of the wage determination is attached as a supporting document.')
    return doc


# ─── 04b_Proposal_Technical.docx ─────────────────────────────────────────────

def build_04b():
    doc = Document()
    add_title(doc, 'TECHNICAL PROPOSAL')
    add_p(doc, '{{SOL_TITLE}}', bold=True)
    add_p(doc, 'Solicitation {{SOL_NUMBER}}')
    add_p(doc, 'Submitted by: {{ENTITY_NAME}}')
    add_p(doc, 'Date: {{TODAY_DATE}}')

    doc.add_paragraph()
    add_h2(doc, '1. Executive Summary')
    add_p(doc, '{{ENTITY_NAME}} is pleased to submit this technical proposal in response to {{SOL_NUMBER}}. {{ENTITY_NAME}} understands the requirement, has the personnel and processes to perform, and offers a compliant, efficient approach that meets every Section L submission instruction and every Section M evaluation factor.')

    add_h2(doc, '2. Understanding of the Requirement')
    add_p(doc, '{{SOL_SCOPE_SUMMARY}}')
    doc.add_paragraph()
    add_p(doc, 'Place of performance: {{SOL_PLACE_OF_PERFORMANCE}}')
    add_p(doc, 'Period of performance: {{SOL_BASE_PERIOD}}, with {{SOL_OPTION_YEARS}}.')
    add_p(doc, 'Work locations include:')
    add_p(doc, '{{SOL_WORK_LOCATIONS}}')

    add_h2(doc, '3. Technical Approach')
    add_h3(doc, '3.1 Service Delivery Model')
    add_p(doc, '{{TECHNICAL_SERVICE_DELIVERY_NARRATIVE}}')
    add_h3(doc, '3.2 Staffing Plan')
    add_p(doc, '{{TECHNICAL_STAFFING_NARRATIVE}}')
    add_h3(doc, '3.3 Equipment, Supplies, and Materials')
    add_p(doc, '{{TECHNICAL_EQUIPMENT_NARRATIVE}}')
    add_h3(doc, '3.4 Schedule and Sequencing')
    add_p(doc, '{{TECHNICAL_SCHEDULE_NARRATIVE}}')

    add_h2(doc, '4. Management Approach')
    add_p(doc, '{{ENTITY_NAME}} will serve as prime contractor under {{ENTITY_PROPOSAL_LEAD}}, who will be the single point of contact for the Contracting Officer and the Contracting Officer\'s Representative. The management approach is built around clear lines of authority, weekly performance reporting, and proactive problem resolution.')
    add_p(doc, '{{TECHNICAL_MANAGEMENT_NARRATIVE}}')

    add_h2(doc, '5. Quality Control Plan')
    add_p(doc, '{{TECHNICAL_QC_NARRATIVE}}')

    add_h2(doc, '6. Safety Plan')
    add_p(doc, '{{TECHNICAL_SAFETY_NARRATIVE}}')

    add_h2(doc, '7. Past Performance')
    add_p(doc, '{{TECHNICAL_PAST_PERFORMANCE_NARRATIVE}}')

    add_h2(doc, '8. Subcontractor Approach')
    add_p(doc, '{{TECHNICAL_SUBCONTRACTOR_NARRATIVE}}')

    add_h2(doc, '9. Transition Plan')
    add_p(doc, '{{TECHNICAL_TRANSITION_NARRATIVE}}')

    add_h2(doc, '10. Compliance Cross-Walk to Sections L and M')
    add_p(doc, 'The following table maps each Section L submission requirement and Section M evaluation factor to the location in this proposal where it is addressed.')
    doc.add_paragraph()
    add_p(doc, '{{SECTION_L_M_CROSSWALK_TABLE}}')

    add_h2(doc, '11. Certification')
    add_p(doc, '{{ENTITY_NAME}} certifies that this technical proposal is accurate, complete, and made in good faith. {{ENTITY_NAME}} accepts the terms of {{SOL_NUMBER}} and is prepared to perform in accordance with the resulting contract.')
    doc.add_paragraph()
    add_p(doc, 'Signed: ____________________________________   Date: ____________________')
    add_p(doc, '{{ENTITY_PROPOSAL_LEAD}}, {{ENTITY_PROPOSAL_LEAD_TITLE}}')
    add_p(doc, '{{ENTITY_NAME}}')
    add_p(doc, 'UEI: {{ENTITY_UEI}} | CAGE: {{ENTITY_CAGE}}')
    return doc


# ─── 06_Sub_Outreach_Email.docx ──────────────────────────────────────────────

def build_06():
    """SANITIZED — OPSEC. NO sol#, agency, CO. Generic phrasing only."""
    doc = Document()
    add_title(doc, 'COLD OUTREACH EMAIL TO SUBCONTRACTOR', size=14)
    add_p(doc, 'Internal-use template. Renderer fills sub-specific tokens; OPSEC enforces no agency leak.', bold=True)
    doc.add_paragraph()

    add_h2(doc, 'Email Subject Line')
    add_p(doc, 'Subject: Federal contract opportunity — {{SUB_SOW_TITLE}} in {{SOL_GENERIC_LOCATION}} — pricing requested by {{RETURN_BY_DATE_FRIENDLY}}')

    doc.add_paragraph()
    add_h2(doc, 'Email Body')
    add_p(doc, 'Hi {{SUB_FIRST_NAME}},')
    doc.add_paragraph()
    add_p(doc, 'My name is {{ENTITY_PROPOSAL_LEAD}} and I lead proposals at {{ENTITY_NAME}}. We are preparing a bid on a federal contract for {{SUB_SOW_TITLE}} services at {{SOL_GENERIC_LOCATION}}.')
    doc.add_paragraph()
    add_p(doc, 'The work covers {{SOL_SCOPE_GENERIC_DESCRIPTION}}. The performance period runs {{SOL_BASE_PERIOD_FRIENDLY}}, with option years that could extend the work through {{SOL_FINAL_PERFORMANCE_END_DATE}}.')
    doc.add_paragraph()
    add_p(doc, 'I am reaching out because your company is well positioned for the field work in this region, and we would like to include your pricing in our proposal.')
    doc.add_paragraph()
    add_p(doc, 'If you are interested, please reply to this email and I will send you:')
    add_bullet(doc, 'A scope of work summary written for subcontractors (no agency identifiers)')
    add_bullet(doc, 'A pre-filled Subcontractor Information & Pricing Form (SIF) so you only fill the gaps')
    add_bullet(doc, 'The pricing return date — pricing is due back to us by {{RETURN_BY_DATE_FRIENDLY}}')
    doc.add_paragraph()
    add_p(doc, 'A few quick questions to make sure this is a fit:')
    add_bullet(doc, 'Is your team available to begin work on or around {{SOL_PERFORMANCE_START_DATE}}?')
    add_bullet(doc, 'Does your insurance currently meet or exceed federal contract minimums (general liability, auto, workers\' comp)?')
    add_bullet(doc, 'Do you hold the licenses or certifications typical for {{SUB_SOW_TITLE}} work in {{SOL_LOCATION_STATE}}?')
    doc.add_paragraph()
    add_p(doc, 'You can reach me at {{ENTITY_PHONE}} or reply directly to this email at {{ENTITY_PUBLIC_EMAIL}}. If you would prefer a quick call to discuss, send me two or three time windows that work this week and I\'ll confirm one.')
    doc.add_paragraph()
    add_p(doc, 'Thank you for your time. I look forward to hearing back.')
    doc.add_paragraph()
    add_p(doc, 'Regards,')
    add_p(doc, '{{ENTITY_PROPOSAL_LEAD}}')
    add_p(doc, '{{ENTITY_PROPOSAL_LEAD_TITLE}}')
    add_p(doc, '{{ENTITY_NAME}}')
    add_p(doc, '{{ENTITY_PHONE}} | {{ENTITY_PUBLIC_EMAIL}}')

    doc.add_paragraph()
    add_h2(doc, 'OPSEC Checklist (DELETE BEFORE SENDING)')
    add_bullet(doc, 'Email body contains NO solicitation number')
    add_bullet(doc, 'Email body contains NO agency name')
    add_bullet(doc, 'Email body contains NO sub-agency or contracting officer')
    add_bullet(doc, 'Email body contains NO SAM.gov URL or notice ID')
    add_bullet(doc, 'Email body contains NO estimated contract value')
    add_bullet(doc, 'Email body contains NO incumbent or USASpending data')
    add_bullet(doc, 'Subject line contains NO solicitation number or agency')

    return doc


# ─── 09_To_Do_List.docx ──────────────────────────────────────────────────────

def build_09():
    doc = Document()
    add_title(doc, 'BID EXECUTION TO-DO LIST')
    add_p(doc, '{{SOL_TITLE}}', bold=True)
    add_p(doc, 'Solicitation {{SOL_NUMBER}} | Due {{SOL_RESPONSE_DEADLINE_FRIENDLY}}')
    add_p(doc, 'Owner: {{ENTITY_PROPOSAL_LEAD}} | Compiled: {{TODAY_DATE}}')
    doc.add_paragraph()

    sections = [
        ('Day 0 — Intake (do today)', [
            'Read full solicitation cover-to-cover; flag every Section L and Section M item',
            'Pull amendments from sam.gov; date-stamp the latest amendment',
            'Confirm SAM.gov registration is active and Reps & Certs are current',
            'Confirm the assigned NAICS code is on our SAM profile and we meet the size standard',
            'Decision: bid / no-bid (record rationale in the bid_decision_notes field on the gov_lead)',
        ]),
        ('Days 1-3 — Compliance Matrix', [
            'Build Section L / Section M cross-walk in the proposal_compliance_items table',
            'List required attachments and forms (SF-1449, Reps & Certs, SPRS, etc.)',
            'List required certifications (e.g. {{SOL_REQUIRED_LICENSE_NAME}})',
            'List required insurance levels and confirm we meet them',
            'Confirm wage determination {{WAGE_DETERMINATION_NUMBER}} is the version cited in the solicitation',
        ]),
        ('Days 2-5 — Subcontractor Outreach', [
            'Generate Subcontractor_Search.docx via the Bid Starter package',
            'Verify each candidate sub via SAM.gov (UEI active, set-asides, NAICS match)',
            'For top candidates, generate per-sub SIF + outreach email (Phase 2)',
            'Send outreach; track responses in proposal_compliance_items',
            'Lock in 1-2 subs for the bid by {{RETURN_BY_DATE_FRIENDLY}}',
        ]),
        ('Days 3-8 — Pricing', [
            'Pull historical award data from USASpending for similar NAICS / agency',
            'Build CLIN pricing in Government Pricing Calculator',
            'Validate sub pricing against received SIFs',
            'Run pricing math audit (qty x unit = extended; sums; no negatives)',
            'Document margin and overhead assumptions in 04a_Proposal_Pricing.docx',
        ]),
        ('Days 5-10 — Technical Narrative', [
            'Draft 04b_Proposal_Technical.docx using the entity boilerplate (safety, QC, mgmt)',
            'Tailor every section to specific solicitation requirements (no copy-paste)',
            'Run drafted narrative through deterministic humanizer',
            'Confirm Section L compliance: every required topic addressed',
            'Confirm Section M compliance: every evaluation factor scored',
        ]),
        ('Days 8-12 — Validation & Pink Team', [
            'Run /api/proposals/validate (5-pass deterministic validator)',
            'Run /api/proposals/pink-team-review (adversarial reviewer)',
            'Apply safe fixes; resolve every red and yellow finding',
            'Re-run validator until all-green',
            'Confirm address protocol: federal docs use Wild Orchid; non-federal use Culpeper',
        ]),
        ('Day before due — Final Submission', [
            'Convert any DOCX deliverables to PDF where solicitation requires',
            'Sign required forms (SF-1449 Block 30a, etc.)',
            'Verify file names match required convention exactly',
            'Verify total file size is under any submission portal limit',
            'Run 13_Submission_Checklist.docx end-to-end',
        ]),
        ('Submission day', [
            'Send / upload at least 2 hours before deadline',
            'Capture submission confirmation (email receipt or portal screenshot) into proposal record',
            'Mark proposal status submitted in mog-tracker',
            'Schedule retro for the day after award decision',
        ]),
    ]
    for header, items in sections:
        add_h2(doc, header)
        for item in items:
            add_bullet(doc, item)
        doc.add_paragraph()

    add_h2(doc, 'Notes and Open Items')
    add_p(doc, '{{TODO_OPEN_ITEMS_BLOCK}}')
    return doc


# ─── 10_Risk_Log.docx ────────────────────────────────────────────────────────

def build_10():
    doc = Document()
    add_title(doc, 'RISK LOG')
    add_p(doc, '{{SOL_TITLE}}', bold=True)
    add_p(doc, 'Solicitation {{SOL_NUMBER}}')
    add_p(doc, 'Owner: {{ENTITY_PROPOSAL_LEAD}} | Updated: {{TODAY_DATE}}')

    doc.add_paragraph()
    add_h2(doc, 'Scoring Key')
    add_bullet(doc, 'Likelihood: 1 (rare) - 2 (unlikely) - 3 (possible) - 4 (likely) - 5 (almost certain)')
    add_bullet(doc, 'Impact: 1 (negligible) - 2 (minor) - 3 (moderate) - 4 (major) - 5 (catastrophic)')
    add_bullet(doc, 'Risk score = Likelihood x Impact. >= 12 requires written mitigation plan.')

    doc.add_paragraph()
    add_h2(doc, 'Risk Register')
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    headers = ['#', 'Risk', 'Likelihood', 'Impact', 'Score', 'Mitigation', 'Owner']
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True

    # Common risks for any federal services bid (template starter)
    common_risks = [
        ('Subcontractor backs out before award', '3', '4', '12',
         'Identify 2 backup subs in 07_Subcontractor_Search; secure LOI before submission',
         '{{ENTITY_PROPOSAL_LEAD}}'),
        ('SCA wage rates change between submission and option year exercise', '3', '3', '9',
         'Cite FAR 52.222-43 for equitable adjustment; price option years with 3-5% inflation cushion',
         '{{ENTITY_PROPOSAL_LEAD}}'),
        ('Solicitation amendment lands within 7 days of due date', '3', '4', '12',
         'Daily sam.gov amendment check; pre-allocate 24 hours buffer for last-minute revisions',
         '{{ENTITY_PROPOSAL_LEAD}}'),
        ('Required license or insurance certificate not in hand by due date', '2', '5', '10',
         'List required docs in compliance matrix on Day 1; chase weekly until confirmed',
         '{{ENTITY_PROPOSAL_LEAD}}'),
        ('LPTA pricing pressure forces unsustainable margins', '4', '3', '12',
         'Set walk-away price floor before bid; document staffing model defends the price',
         '{{ENTITY_PROPOSAL_LEAD}}'),
        ('Site visit conflicts with day job or family obligation', '2', '2', '4',
         'Assign sub or local partner to attend; record visit notes within 24 hours',
         '{{ENTITY_PROPOSAL_LEAD}}'),
        ('Submission portal (PIEE, sam.gov) outage on due day', '2', '5', '10',
         'Submit at least 24 hours early; maintain email backup with CO and contract specialist',
         '{{ENTITY_PROPOSAL_LEAD}}'),
        ('OPSEC leak in sub-facing docs (sol#, agency, CO)', '2', '4', '8',
         'Generator enforces sanitization; manual final scan before any sub email',
         '{{ENTITY_PROPOSAL_LEAD}}'),
    ]
    for i, row in enumerate(common_risks, 1):
        table.add_row()
        c = table.cell(i, 0)
        c.text = str(i)
        for j, val in enumerate(row):
            table.cell(i, j+1).text = val

    doc.add_paragraph()
    add_p(doc, '{{ADDITIONAL_RISKS_BLOCK}}')

    add_h2(doc, 'Risks Closed')
    add_p(doc, 'Document risks that have been retired with date and how they were retired.')
    add_p(doc, '{{CLOSED_RISKS_BLOCK}}')

    return doc


# ─── 12_Game_Plan.docx ───────────────────────────────────────────────────────

def build_12():
    doc = Document()
    add_title(doc, 'GAME PLAN')
    add_p(doc, '{{SOL_TITLE}}', bold=True)
    add_p(doc, 'Solicitation {{SOL_NUMBER}} | Due {{SOL_RESPONSE_DEADLINE_FRIENDLY}}')
    add_p(doc, 'Author: {{ENTITY_PROPOSAL_LEAD}} | Compiled: {{TODAY_DATE}}')

    doc.add_paragraph()
    add_h2(doc, '1. Why We Are Bidding')
    add_p(doc, '{{GAME_PLAN_RATIONALE}}')

    add_h2(doc, '2. Win Probability Assessment')
    add_kv_table(doc, [
        ('Win probability:', '{{WIN_PROBABILITY_RANGE}}'),
        ('Confidence level:', '{{WIN_PROBABILITY_CONFIDENCE}}'),
        ('Primary win driver:', '{{PRIMARY_WIN_DRIVER}}'),
        ('Primary risk:', '{{PRIMARY_LOSS_RISK}}'),
    ])

    add_h2(doc, '3. Competitive Positioning')
    add_p(doc, '{{COMPETITIVE_POSITIONING_NARRATIVE}}')

    add_h2(doc, '4. Pricing Strategy')
    add_kv_table(doc, [
        ('Target price range:', '{{TARGET_PRICE_RANGE}}'),
        ('Walk-away floor:', '{{WALK_AWAY_PRICE}}'),
        ('Margin assumption:', '{{MARGIN_ASSUMPTION}}'),
        ('Pricing approach:', '{{PRICING_APPROACH}}'),
    ])

    add_h2(doc, '5. Teaming Strategy')
    add_p(doc, '{{TEAMING_STRATEGY_NARRATIVE}}')
    add_kv_table(doc, [
        ('Primary teaming partner:', '{{TEAMING_PARTNER_NAME}}'),
        ('Partner role:', '{{TEAMING_PARTNER_ROLE}}'),
        ('Backup partner:', '{{TEAMING_PARTNER_BACKUP}}'),
        ('LOI status:', '{{TEAMING_LOI_STATUS}}'),
    ])

    add_h2(doc, '6. Differentiators')
    add_bullet(doc, '{{DIFFERENTIATOR_1}}')
    add_bullet(doc, '{{DIFFERENTIATOR_2}}')
    add_bullet(doc, '{{DIFFERENTIATOR_3}}')

    add_h2(doc, '7. Top 3 Risks and Mitigations')
    add_p(doc, 'Refer to 10_Risk_Log.docx for the complete register. The top three risks for this bid:')
    add_bullet(doc, '{{TOP_RISK_1}}')
    add_bullet(doc, '{{TOP_RISK_2}}')
    add_bullet(doc, '{{TOP_RISK_3}}')

    add_h2(doc, '8. Decision Points')
    add_p(doc, 'These are the moments at which we will re-evaluate go/no-go:')
    add_bullet(doc, 'After amendment review — confirm scope has not changed materially')
    add_bullet(doc, 'After sub pricing returns — confirm we can hit our target price')
    add_bullet(doc, 'After validator pink-team — confirm no fatal flaws')
    add_bullet(doc, 'Submission day - 24h — final go/no-go on whether to submit at all')

    add_h2(doc, '9. After-Action')
    add_p(doc, 'Whether we win or lose, we capture the result in the proposal_retros table within 7 days of the award decision. Lessons learned feed the next bid.')

    return doc


# ─── 08_Tracker_Workbook.xlsx (Excel) ────────────────────────────────────────

def build_08_xlsx():
    wb = openpyxl.Workbook()

    # Default first sheet -> Overview
    ws = wb.active
    ws.title = 'Overview'
    overview_rows = [
        ('Field', 'Value'),
        ('Solicitation Number', '{{SOL_NUMBER}}'),
        ('Solicitation Title', '{{SOL_TITLE}}'),
        ('Agency', '{{SOL_AGENCY}}'),
        ('Sub-Agency', '{{SOL_SUB_AGENCY}}'),
        ('NAICS', '{{SOL_NAICS}} - {{SOL_NAICS_DESCRIPTION}}'),
        ('Set-Aside', '{{SOL_SET_ASIDE}}'),
        ('Contract Type', '{{SOL_CONTRACT_TYPE}}'),
        ('Place of Performance', '{{SOL_PLACE_OF_PERFORMANCE}}'),
        ('Base Period', '{{SOL_BASE_PERIOD}}'),
        ('Option Years', '{{SOL_OPTION_YEARS}}'),
        ('Response Deadline', '{{SOL_RESPONSE_DEADLINE}}'),
        ('Estimated Value', '{{SOL_ESTIMATED_VALUE}}'),
        ('Proposal Lead', '{{ENTITY_PROPOSAL_LEAD}}'),
        ('Today', '{{TODAY_DATE}}'),
    ]
    for r in overview_rows:
        ws.append(r)
    for c in range(1, 3):
        ws.cell(row=1, column=c).font = Font(bold=True)
    ws.column_dimensions['A'].width = 28
    ws.column_dimensions['B'].width = 60

    # CLIN Pricing sheet
    ws2 = wb.create_sheet('CLIN_Pricing')
    headers = ['CLIN', 'Description', 'Period', 'Qty', 'Unit', 'Unit Price', 'Extended Price', 'Notes']
    ws2.append(headers)
    ws2.append(['{{SOL_CLIN_TABLE}}', '', '', '', '', '', '', ''])
    for c in range(1, len(headers)+1):
        ws2.cell(row=1, column=c).font = Font(bold=True)
    for col, w in zip('ABCDEFGH', [10, 35, 22, 8, 8, 14, 16, 22]):
        ws2.column_dimensions[col].width = w

    # Subcontractors sheet
    ws3 = wb.create_sheet('Subcontractors')
    sub_headers = ['Company', 'UEI', 'CAGE', 'Contact', 'Phone', 'Email', 'NAICS Match',
                   'Cert', 'Verification', 'Score', 'Source', 'LOI Status', 'Notes']
    ws3.append(sub_headers)
    ws3.append(['{{SUB_CANDIDATES_BLOCK}}'] + [''] * (len(sub_headers)-1))
    for c in range(1, len(sub_headers)+1):
        ws3.cell(row=1, column=c).font = Font(bold=True)

    # Compliance sheet
    ws4 = wb.create_sheet('Compliance')
    comp_headers = ['Section', 'Item', 'Required Doc / Action', 'Owner', 'Due', 'Status', 'File / Evidence']
    ws4.append(comp_headers)
    ws4.append(['{{COMPLIANCE_MATRIX_BLOCK}}'] + [''] * (len(comp_headers)-1))
    for c in range(1, len(comp_headers)+1):
        ws4.cell(row=1, column=c).font = Font(bold=True)

    # Risks sheet
    ws5 = wb.create_sheet('Risks')
    risk_headers = ['#', 'Risk', 'Likelihood (1-5)', 'Impact (1-5)', 'Score', 'Mitigation', 'Owner', 'Status']
    ws5.append(risk_headers)
    ws5.append(['{{ADDITIONAL_RISKS_BLOCK}}'] + [''] * (len(risk_headers)-1))
    for c in range(1, len(risk_headers)+1):
        ws5.cell(row=1, column=c).font = Font(bold=True)

    # Pricing summary sheet
    ws6 = wb.create_sheet('Pricing_Summary')
    summary = [
        ('Component', 'Amount', 'Notes'),
        ('Direct labor (loaded)', '', ''),
        ('Subcontractor cost', '', ''),
        ('Materials and supplies', '', ''),
        ('Equipment and vehicle', '', ''),
        ('Travel', '', ''),
        ('Overhead', '', ''),
        ('G&A', '', ''),
        ('Profit / fee', '', ''),
        ('Total proposed price', '', ''),
        ('Walk-away floor', '{{WALK_AWAY_PRICE}}', ''),
        ('Target range', '{{TARGET_PRICE_RANGE}}', ''),
    ]
    for r in summary:
        ws6.append(r)
    for c in range(1, 4):
        ws6.cell(row=1, column=c).font = Font(bold=True)
    ws6.column_dimensions['A'].width = 30
    ws6.column_dimensions['B'].width = 18
    ws6.column_dimensions['C'].width = 40

    return wb


# ─── Save all ────────────────────────────────────────────────────────────────

build_03().save(os.path.join(OUT, '03_Wage_Determination.docx'))
print('Built: 03_Wage_Determination.docx')
build_04b().save(os.path.join(OUT, '04b_Proposal_Technical.docx'))
print('Built: 04b_Proposal_Technical.docx')
build_06().save(os.path.join(OUT, '06_Sub_Outreach_Email.docx'))
print('Built: 06_Sub_Outreach_Email.docx')
build_09().save(os.path.join(OUT, '09_To_Do_List.docx'))
print('Built: 09_To_Do_List.docx')
build_10().save(os.path.join(OUT, '10_Risk_Log.docx'))
print('Built: 10_Risk_Log.docx')
build_12().save(os.path.join(OUT, '12_Game_Plan.docx'))
print('Built: 12_Game_Plan.docx')
build_08_xlsx().save(os.path.join(OUT, '08_Tracker_Workbook.xlsx'))
print('Built: 08_Tracker_Workbook.xlsx')
