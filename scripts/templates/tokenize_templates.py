#!/usr/bin/env python3
"""
Tokenize 12 bid-level templates from src/lib/bid-templates/source/
Output to src/lib/bid-templates/tokenized/

Approach: per-file substitution map (literal source -> token).
Walk every paragraph and every table cell, apply substitutions.
Verify each output file has >0 unique {{TOKEN}} placeholders.
"""

import os, sys, re, json, shutil, subprocess
from copy import deepcopy
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = '/sessions/beautiful-lucid-ritchie/mnt/mog-tracker-app/src/lib/bid-templates'
SRC = os.path.join(ROOT, 'source')
OUT = os.path.join(ROOT, 'tokenized')

os.makedirs(OUT, exist_ok=True)

# ---------------------------------------------------------------------------
# Substitution helpers
# ---------------------------------------------------------------------------

def replace_in_paragraph(p, subs):
    """Apply ordered (find, replace) substitutions to a paragraph.
    Operates on combined paragraph text and rebuilds runs as a single run
    preserving the first run's style. This loses inline formatting variations
    (bold/italic mid-paragraph) but reliably handles run-spanning matches.
    """
    full = ''.join(r.text for r in p.runs)
    if not full:
        return False
    new_full = full
    for find, repl in subs:
        if find and find in new_full:
            new_full = new_full.replace(find, repl)
    if new_full == full:
        return False
    # Preserve first run's properties; clear others
    if p.runs:
        first_run = p.runs[0]
        for r in p.runs[1:]:
            r.text = ''
        first_run.text = new_full
    else:
        p.text = new_full
    return True


def replace_in_cell(cell, subs):
    changed = False
    for p in cell.paragraphs:
        if replace_in_paragraph(p, subs):
            changed = True
    return changed


def apply_subs_to_doc(doc, subs):
    """Apply substitutions to all paragraphs in body, tables, headers, footers."""
    # Body paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p, subs)
    # Tables (recursively for nested tables)
    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    replace_in_cell(cell, subs)
                    if cell.tables:
                        walk_tables(cell.tables)
    walk_tables(doc.tables)
    # Headers and footers
    for section in doc.sections:
        for hdr_or_ftr in [section.header, section.footer,
                           section.first_page_header, section.first_page_footer,
                           section.even_page_header, section.even_page_footer]:
            for p in hdr_or_ftr.paragraphs:
                replace_in_paragraph(p, subs)
            for t in hdr_or_ftr.tables:
                for row in t.rows:
                    for cell in row.cells:
                        replace_in_cell(cell, subs)


def replace_table_with_placeholder(doc, table_index, placeholder_text):
    """Remove a table by index and insert a placeholder paragraph in its position."""
    if table_index >= len(doc.tables):
        return False
    table = doc.tables[table_index]
    tbl_elem = table._tbl
    # Insert paragraph before
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = placeholder_text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    tbl_elem.addprevious(new_p)
    tbl_elem.getparent().remove(tbl_elem)
    return True


def count_tokens_in_doc(path):
    d = Document(path)
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    parts.append(cp.text)
    full = ' '.join(parts)
    tokens = re.findall(r'\{\{[A-Z_]+\}\}', full)
    unique = sorted(set(tokens))
    return tokens, unique


# ---------------------------------------------------------------------------
# Per-file substitution maps
# ---------------------------------------------------------------------------

# Common substitutions across ALL files (entity identifiers, dates, etc.)
# Order matters: longer strings first to avoid partial-match collisions.
COMMON_SUBS = [
    # --- Entity identifiers (must come BEFORE plain "Exousia") ---
    ('Exousia Solutions LLC', '{{ENTITY_NAME}}'),
    ('EXOUSIA SOLUTIONS LLC', '{{ENTITY_NAME_UPPER}}'),
    ('XNZ2KYQYK566', '{{ENTITY_UEI}}'),
    ('0ENQ3', '{{ENTITY_CAGE}}'),
    ('admin@exousiaofficial.com', '{{ENTITY_SAM_EMAIL}}'),
    ('Info@exousias.com', '{{ENTITY_PUBLIC_EMAIL}}'),
    ('info@exousias.com', '{{ENTITY_PUBLIC_EMAIL}}'),
    ('(571) 556-2227', '{{ENTITY_PHONE}}'),
    ('(571) 622-9133', '{{ENTITY_PHONE}}'),
    ('Emmanuela Wireko-Brobbey', '{{ENTITY_PROPOSAL_LEAD}}'),
]

# Per-file additional substitutions
FILE_SUBS = {
    '01_Contract_Intel_Sheet.docx': [
        # Sol identifiers
        ('W912PM26QA014', '{{SOL_NUMBER}}'),
        ("Custodial and Refuse Collection Services at Poe's Ridge Recreation Area",
         '{{SOL_TITLE}}'),
        ('Department of Defense, US Army Corps of Engineers (USACE), Wilmington District',
         '{{SOL_AGENCY}}'),
        ('561720 - Janitorial Services', '{{SOL_NAICS}} - {{SOL_NAICS_DESCRIPTION}}'),
        ('Small Business 100% Set-Aside', '{{SOL_SET_ASIDE}}'),
        ('LPTA (Lowest Price Technically Acceptable)', '{{SOL_EVALUATION_METHOD}}'),
        ('09 APR 2026 at 2:00 PM Eastern Time', '{{SOL_RESPONSE_DEADLINE}}'),
        ('Electronic submission in Adobe PDF format via email', '{{SOL_SUBMISSION_METHOD}}'),
        ('W912PM26QA014_COMPANY NAME_PRICING and W912PM26QA014_COMPANY NAME_TECHNICAL',
         '{{SOL_FILE_NAMING}}'),
        # Contracting officer
        ('Benjamin Rickman', '{{CO_NAME}}'),
        ('910-251-4494', '{{CO_PHONE}}'),
        ('benjamin.t.rickman@usace.army.mil', '{{CO_EMAIL}}'),
        ('Troy D. Small', '{{KO_CONTACT_SPECIALIST_NAME}}'),
        ('troy.d.small@usace.army.mil', '{{KO_CONTACT_SPECIALIST_EMAIL}}'),
        # Locations
        ("Poe's Ridge Recreation Area, Moncure, NC 27559", '{{SOL_PLACE_OF_PERFORMANCE}}'),
        ('B. Everett Jordan Dam VAC, 2080 Jordan Dam Rd', '{{SOL_SITE_1}}'),
        ('Jordan Dam Tailrace Fishing Area, 2060 Jordan Dam Rd', '{{SOL_SITE_2}}'),
        ("Poe's Ridge Boat Ramp, 935 Jordan Dam Rd", '{{SOL_SITE_3}}'),
        # Period of performance
        ('01 JUN 2026 to 28 FEB 2027 (9 months)', '{{SOL_BASE_PERIOD}}'),
        ('4 option years (12 months each) through 28 FEB 2031', '{{SOL_OPTION_YEARS}}'),
        # Wage determination
        ('WD 2015-4373 Rev 32, Chatham County NC',
         '{{WAGE_DETERMINATION_NUMBER}}, {{WAGE_DETERMINATION_LOCATION}}'),
        ('Janitor: $16.16/hr + $5.17 fringe', '{{WAGE_EQUIVALENT_RATES}}'),
        ('Janitor: $17.14/hr\nHousekeeping Aide: $17.14/hr\nLaborer Grounds Maintenance: $18.07/hr\nWindow Cleaner: $18.63/hr',
         '{{WAGE_KEY_RATES}}'),
        # CLINs
        ('Base 9 months FFP (01 JUN 2026 - 28 FEB 2027)', '{{CLIN_0001_DESCRIPTION}}'),
        ('Option Year 1 (12 months)', '{{CLIN_OY1_DESCRIPTION}}'),
        ('Option Year 2 (12 months)', '{{CLIN_OY2_DESCRIPTION}}'),
        ('Option Year 3 (12 months)', '{{CLIN_OY3_DESCRIPTION}}'),
        ('Option Year 4 (12 months) through 28 FEB 2031', '{{CLIN_OY4_DESCRIPTION}}'),
        # Research items - keep [PENDING] markers as tokens
        ('[PENDING - USASpending.gov lookup required]', '{{SOL_RESEARCH_PENDING}}'),
        ('[PENDING - SAM.gov amendment check required]', '{{SOL_AMENDMENT_PENDING}}'),
    ],

    '02_USASpending_Deep_Dive.docx': [
        ('W912PM26QA014', '{{SOL_NUMBER}}'),
        ('USACE Wilmington District', '{{SOL_SUB_AGENCY}}'),
        ('U.S. Army Corps of Engineers, Wilmington District (SAW)', '{{SOL_AGENCY}}'),
        ('USACE Wilmington', '{{SOL_SUB_AGENCY}}'),
        ('USACE', '{{SOL_AGENCY_ABBREV}}'),
        ('BEJ Custodial Contract', '{{SOL_TITLE}}'),
        ("Poe's Ridge Recreation Area custodial services contract", '{{SOL_TITLE}}'),
        ("Poe's Ridge Recreation Area", '{{SOL_PLACE_OF_PERFORMANCE}}'),
        ('Moncure', '{{SOL_LOCATION_CITY}}'),
        ('B. Everett Jordan Lake', '{{SOL_FACILITY_NAME}}'),
        ('North Carolina', '{{SOL_LOCATION_STATE}}'),
        ('NC', '{{SOL_LOCATION_STATE_ABBREV}}'),
        ('WD 2015-4373 Rev 32', '{{WAGE_DETERMINATION_NUMBER}}'),
        ('561720', '{{SOL_NAICS}}'),
        ('$180K-$220K', '{{ESTIMATE_HIGH_RANGE}}'),
        ('$140K-$170K', '{{ESTIMATE_MID_RANGE}}'),
        ('$145K-$160K', '{{TARGET_PRICE_RANGE}}'),
        ('$145K-160K', '{{TARGET_PRICE_RANGE}}'),
        ('$17.14/hr', '{{WAGE_BASE_JANITOR}}'),
        ('$18.07/hr', '{{WAGE_BASE_LABORER}}'),
        ('$18.63/hr', '{{WAGE_BASE_WINDOW}}'),
        ('$16.16/hr', '{{WAGE_EQUIV_JANITOR}}'),
        ('$5.55/hr', '{{WAGE_FRINGE}}'),
        ('$5.09/hr', '{{WAGE_FRINGE_EO13706}}'),
        ('$5.17 fringe', '{{WAGE_EQUIV_FRINGE}}'),
        ('55-65%', '{{WIN_PROBABILITY_RANGE}}'),
        ('Ironhouse Janitorial & Landscaping', '{{TEAMING_PARTNER_NAME}}'),
        ('Ironhouse', '{{TEAMING_PARTNER_SHORT}}'),
        ('SEC via Zvolvant/GDIT', '{{ENTITY_PAST_PERFORMANCE_NOTE}}'),
        ('SEC compliance', '{{ENTITY_PAST_PERFORMANCE_AREA}}'),
        ('VA-based', '{{TEAMING_PARTNER_LOCATION}}'),
    ],

    '04_Proposal_Pricing.docx': [
        ('W912PM26QA014', '{{SOL_NUMBER}}'),
        ('Custodial and Refuse Collection Services', '{{SOL_TITLE}}'),
        ("Poe's Ridge Recreation Area, Moncure, North Carolina", '{{SOL_PLACE_OF_PERFORMANCE}}'),
        ('Poe’s Ridge Recreation Area, Moncure, North Carolina', '{{SOL_PLACE_OF_PERFORMANCE}}'),
        ('Moncure, NC', '{{SOL_LOCATION_CITY_STATE}}'),
        ('Chatham County, NC', '{{WAGE_DETERMINATION_LOCATION}}'),
        ('WD 2015-4373 Rev 32', '{{WAGE_DETERMINATION_NUMBER}}'),
        ('USACE', '{{SOL_AGENCY_ABBREV}}'),
        ('[PROPOSAL DATE]', '{{TODAY_DATE}}'),
        ('UEI: XNZ2KYQYK566 | CAGE: 0ENQ3', 'UEI: {{ENTITY_UEI}} | CAGE: {{ENTITY_CAGE}}'),
        ('Email: admin@exousiaofficial.com | Phone: (571) 556-2227',
         'Email: {{ENTITY_SAM_EMAIL}} | Phone: {{ENTITY_PHONE}}'),
        ('June 1, 2026', '{{SOL_PERFORMANCE_START_DATE}}'),
        ('$17.14', '{{WAGE_BASE_JANITOR_RATE}}'),
        ('$18.07', '{{WAGE_BASE_LABORER_RATE}}'),
        ('$5.55', '{{WAGE_FRINGE_RATE}}'),
        ('$22.69/hr', '{{WAGE_LOADED_JANITOR}}'),
        ('$23.62/hr', '{{WAGE_LOADED_LABORER}}'),
        ('$800-$1,200', '{{SUPPLIES_MONTHLY_RANGE}}'),
        ('Ironhouse', '{{TEAMING_PARTNER_SHORT}}'),
    ],

    '05_Subcontractor_Scope_of_Work.docx': [
        # OPSEC: agency identifiers replaced with generic phrasing (not tokens)
        ('U.S. Army Corps of Engineers', 'the federal customer'),
        ('Army Corps of Engineers', 'the federal customer'),
        ('USACE', 'the federal customer'),
        # Specific federal facility names: replace with generic
        ("Poe's Ridge Recreation Area — Moncure, NC", '{{SOL_GENERIC_LOCATION}}'),
        ('Poe’s Ridge Recreation Area — Moncure, NC', '{{SOL_GENERIC_LOCATION}}'),
        ("Poe's Ridge Recreation Area", '{{SOL_GENERIC_LOCATION}}'),
        ('B. Everett Jordan Dam Visitor Assistance Center (VAC) — 2080 Jordan Dam Rd, Moncure, NC 27559',
         '{{SOL_WORK_SITE_1}}'),
        ('Jordan Dam Tailrace Fishing Area — 2060 Jordan Dam Rd, Moncure, NC 27559',
         '{{SOL_WORK_SITE_2}}'),
        ("Poe's Ridge Boat Ramp — 935 Jordan Dam Rd, Moncure, NC 27559",
         '{{SOL_WORK_SITE_3}}'),
        ('Poe’s Ridge Boat Ramp — 935 Jordan Dam Rd, Moncure, NC 27559',
         '{{SOL_WORK_SITE_3}}'),
        ('B. Everett Jordan Dam project in Chatham County, North Carolina',
         '{{SOL_GENERIC_PROJECT_AREA}}'),
        ('Chatham County, North Carolina', '{{SOL_LOCATION_COUNTY_STATE}}'),
        ('Moncure, NC 27559', '{{SOL_LOCATION_CITY_STATE_ZIP}}'),
        # Dates
        ('June 1, 2026 through February 28, 2027', '{{SOL_BASE_PERIOD_FRIENDLY}}'),
        ('February 28, 2031', '{{SOL_FINAL_PERFORMANCE_END_DATE}}'),
        ('April 5, 2026 (end of day)', '{{RETURN_BY_DATE_FRIENDLY}}'),
    ],

    '07_Subcontractor_Search.docx': [
        # Internal doc - full agency detail OK
        ('W912PM26QA014', '{{SOL_NUMBER}}'),
        ('Moncure, NC (Chatham County)', '{{SOL_PLACE_OF_PERFORMANCE}}'),
        ('Moncure, NC', '{{SOL_LOCATION_CITY_STATE}}'),
        ('Army Corps of Engineers', '{{SOL_AGENCY}}'),
        ('Sanford, 30 min', '{{NEAREST_REGIONAL_HUB_MIN}}'),
        ('Sanford', '{{REGIONAL_HUB_CITY}}'),
        ('North Carolina', '{{SOL_LOCATION_STATE}}'),
        ('June 2026 – February 2027', '{{SOL_BASE_PERIOD_RANGE}}'),
        ('April 5', '{{RETURN_BY_DATE_SHORT}}'),
        ('April 9, 2:00 PM ET', '{{SOL_RESPONSE_DEADLINE_FRIENDLY}}'),
    ],

    '11_Contract_Summary.docx': [
        ('W912PM26QA014', '{{SOL_NUMBER}}'),
        ("Custodial and Refuse Collection Services — Poe's Ridge Recreation Area",
         '{{SOL_TITLE}}'),
        ('Custodial and Refuse Collection Services — Poe’s Ridge Recreation Area',
         '{{SOL_TITLE}}'),
        ('U.S. Army Corps of Engineers, Wilmington District',
         '{{SOL_AGENCY}}, {{SOL_SUB_AGENCY}}'),
        ('April 9, 2026, 2:00 PM Eastern', '{{SOL_RESPONSE_DEADLINE_FRIENDLY}}'),
        ('B. Everett Jordan Dam Visitor Assistance Center — 2080 Jordan Dam Rd, Moncure, NC 27559',
         '{{SOL_SITE_1}}'),
        ('Jordan Dam Tailrace Fishing Area — 2060 Jordan Dam Rd, Moncure, NC 27559',
         '{{SOL_SITE_2}}'),
        ("Poe's Ridge Boat Ramp — 935 Jordan Dam Rd, Moncure, NC 27559",
         '{{SOL_SITE_3}}'),
        ('Poe’s Ridge Boat Ramp — 935 Jordan Dam Rd, Moncure, NC 27559',
         '{{SOL_SITE_3}}'),
        ('Jordan Dam in Moncure, North Carolina', '{{SOL_FACILITY_AND_LOCATION}}'),
        ('Moncure, NC', '{{SOL_LOCATION_CITY_STATE}}'),
        ('June 1, 2026 through February 28, 2027', '{{SOL_BASE_PERIOD_FRIENDLY}}'),
        ('February 28, 2031', '{{SOL_FINAL_PERFORMANCE_END_DATE}}'),
        ('benjamin.t.rickman@usace.army.mil', '{{CO_EMAIL}}'),
    ],

    '13_Submission_Checklist.docx': [
        ('W912BV26QA047', '{{SOL_NUMBER}}'),
        ('April 20, 2026, 2:00 PM CT', '{{SOL_RESPONSE_DEADLINE_FRIENDLY}}'),
        ('19 Mar 2026', '{{SOL_AMENDMENT_CHECK_DATE}}'),
        ('$1,153,850.00', '{{SOL_PRICING_TOTAL_FORMATTED}}'),
        ('$1,153,850', '{{SOL_PRICING_TOTAL_FORMATTED_SHORT}}'),
        ("Sober's KS Pesticide License", '{{SUB_LICENSE_NAME}}'),
        ('Sober’s KS Pesticide License', '{{SUB_LICENSE_NAME}}'),
        ('12/31/2028', '{{SUB_LICENSE_EXPIRY}}'),
        ("Sober's Lawn Care", '{{SUB_NAME}}'),
        ('Sober’s Lawn Care', '{{SUB_NAME}}'),
        ('Sobers_KS_Pesticide_License_278141.pdf', '{{SUB_LICENSE_FILE}}'),
        ('terry.l.hawkins@usace.army.mil', '{{CO_EMAIL}}'),
        ('Wyman.w.walker@usace.army.mil', '{{KO_CONTACT_SPECIALIST_EMAIL}}'),
        ('65 CLINs', '{{SOL_CLIN_COUNT}}'),
    ],

    '14_SIF_Subcontractor_Information_Form.docx': [
        # Header
        ('Mowing and Lawn Maintenance Services | Southeast Kansas',
         '{{SUB_SOW_TITLE}} | {{SOL_GENERIC_LOCATION}}'),
        ('Mowing and Lawn Maintenance Services', '{{SUB_SOW_TITLE}}'),
        ('Southeast Kansas', '{{SOL_GENERIC_LOCATION}}'),
        ('April 2026', '{{TODAY_MONTH_YEAR}}'),
        # Sub-specific
        ('Dear Houston,', 'Dear {{SUB_FIRST_NAME}},'),
        ("Sober's Lawn Care", '{{SUB_COMPANY_NAME}}'),
        ('Sober’s Lawn Care', '{{SUB_COMPANY_NAME}}'),
        ('Houston [Lastname]', '{{SUB_CONTACT_NAME}}'),
        ('(620) 440-2148', '{{SUB_PHONE}}'),
        ('soberslawncare@gmail.com', '{{SUB_EMAIL}}'),
        ('April 18, 2026', '{{RETURN_BY_DATE_FRIENDLY}}'),
        # OPSEC: Bid scope (sub-facing - generic agency)
        ('U.S. Army Corps of Engineers', 'the federal customer'),
        ('two recreational lake projects in southeast Kansas',
         '{{SOL_SCOPE_GENERIC_DESCRIPTION}}'),
        ('Fall River Lake (Greenwood County, KS) and Toronto Lake (Woodson County, KS)',
         '{{SOL_GENERIC_WORK_AREAS}}'),
        ('Fall River Lake', '{{SOL_GENERIC_AREA_1}}'),
        ('Toronto Lake', '{{SOL_GENERIC_AREA_2}}'),
        ('Fall River Lake Office address: 2453 Lake Road, Fall River, KS 67047',
         'Site visit coordination: contact {{ENTITY_PROPOSAL_LEAD}}'),
        # Required documents (bid-specific)
        ("Kansas Commercial Herbicide Applicator's License",
         '{{SOL_REQUIRED_LICENSE_NAME}}'),
        ('Kansas Commercial Herbicide Applicator’s License',
         '{{SOL_REQUIRED_LICENSE_NAME}}'),
        ('Kansas Commercial Herbicide Applicator',
         '{{SOL_REQUIRED_LICENSE_AUTHORITY}}'),
        # NAICS
        ('NAICS 561730 (Landscaping Services, size standard $9.5M)',
         'NAICS {{SOL_NAICS}} ({{SOL_NAICS_DESCRIPTION}}, size standard {{SOL_SIZE_STANDARD}})'),
        # Dates / season
        ('April 1 through September 30', '{{SOL_SEASON_RANGE}}'),
        ('April 2026', '{{SOL_PERFORMANCE_START_MONTH_YEAR}}'),
        # Pricing structure
        ('base year plus four option years (5 years total)',
         '{{SOL_CONTRACT_TERM_DESCRIPTION}}'),
        # Insurance amounts (specific to this bid)
        ('Per Kansas state requirements, minimum $100,000 liability',
         '{{INSURANCE_WORKERS_COMP_REQ}}'),
        ('$500,000 per occurrence for bodily injury; $100,000 for property damage',
         '{{INSURANCE_GL_REQ}}'),
        ('$200,000 per person / $500,000 per occurrence for bodily injury; $20,000 per occurrence for property damage',
         '{{INSURANCE_AUTO_REQ}}'),
        # Times
        ('between 7:00 AM and sunset (or 8:00 PM, whichever is earlier)',
         '{{SOL_WORK_HOURS}}'),
        # Contact entity addresses
        ('Email: admin@exousiaofficial.com', 'Email: {{ENTITY_SAM_EMAIL}}'),
        ('Phone: (571) 622-9133', 'Phone: {{ENTITY_PHONE}}'),
    ],
}

# Tables to replace with block placeholders (file -> [(table_index, placeholder)])
# Note: After we delete tables, indices SHIFT. Use reverse-order processing.
TABLE_REPLACEMENTS = {
    '04_Proposal_Pricing.docx': [
        (0, '{{SOL_CLIN_TABLE}}'),
        (1, '{{LABOR_BREAKDOWN_TABLE}}'),
    ],
    '07_Subcontractor_Search.docx': [
        (0, '{{SUB_CANDIDATES_BLOCK}}'),
        (1, '{{SUB_CANDIDATES_BLOCK_2}}'),
        (2, '{{SUB_CANDIDATES_BLOCK_3}}'),
        (3, '{{SUB_CANDIDATES_BLOCK_4}}'),
    ],
    '02_USASpending_Deep_Dive.docx': [
        (0, '{{WAGE_RATES_TABLE}}'),
        (1, '{{COMPETITIVE_ASSESSMENT_TABLE}}'),
    ],
}

# Deletions (paragraphs to remove entirely from sub-facing docs that leak agency)
# Process by approximate text match.

# ---------------------------------------------------------------------------
# Process each file
# ---------------------------------------------------------------------------

processed = []

for fname in sorted(os.listdir(SRC)):
    if not fname.endswith('.docx'):
        continue
    if not fname.startswith(('01_', '02_', '04_', '05_', '07_', '11_', '13_', '14_')):
        continue

    src_path = os.path.join(SRC, fname)
    out_path = os.path.join(OUT, fname)

    # Copy fresh
    shutil.copy(src_path, out_path)

    doc = Document(out_path)

    # Build full sub list: file-specific first, then common
    file_specific = FILE_SUBS.get(fname, [])
    full_subs = file_specific + COMMON_SUBS

    apply_subs_to_doc(doc, full_subs)

    # Replace tables with block placeholders (process in REVERSE index order)
    if fname in TABLE_REPLACEMENTS:
        for table_idx, placeholder in sorted(TABLE_REPLACEMENTS[fname], reverse=True):
            replace_table_with_placeholder(doc, table_idx, placeholder)

    doc.save(out_path)

    # Verify token count
    tokens, unique = count_tokens_in_doc(out_path)
    processed.append({'file': fname, 'total_tokens': len(tokens), 'unique_tokens': unique})
    print(f'{fname}: {len(tokens)} tokens, {len(unique)} unique')

# ---------------------------------------------------------------------------
# Write audit summary
# ---------------------------------------------------------------------------

print('\n=== SUMMARY ===')
for p in processed:
    if p['total_tokens'] == 0:
        print(f'  FAIL: {p["file"]} has 0 tokens!')
    else:
        print(f'  OK: {p["file"]}')

# Sanity: any file with 0 tokens means the source-string didn't match.
# That'd be a real bug, abort.
if any(p['total_tokens'] == 0 for p in processed):
    print('\nABORTING: some files have 0 tokens. Re-check substitution maps.')
    sys.exit(1)

with open(os.path.join(OUT, '_processing_log.json'), 'w') as f:
    json.dump(processed, f, indent=2)
print('\nProcessing log written.')
